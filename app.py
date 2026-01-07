import io
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import statsmodels.api as sm
import streamlit as st

from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression, LinearRegression, Ridge, Lasso
from sklearn.ensemble import RandomForestClassifier
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.metrics import (
    roc_curve, auc, classification_report, confusion_matrix, ConfusionMatrixDisplay,
    mean_absolute_error, mean_squared_error, r2_score
)
from scipy import stats
from docx import Document


# =========================================================
# Page
# =========================================================
st.set_page_config(
    page_title="Logistic Regression Diagnosis",
    page_icon="📊",
    layout="wide"
)
st.markdown("""
<style>
/* ===== Sidebar menu: button style (no radio circle) ===== */

section[data-testid="stSidebar"]{
  background: linear-gradient(180deg, #0B3A66 0%, #0A2D4E 100%);
}

section[data-testid="stSidebar"] *{
  color: #ffffff !important;
}

/* Hide radio circle */
section[data-testid="stSidebar"] div[role="radiogroup"] input[type="radio"]{
  display: none !important;
}

/* Menu buttons */
section[data-testid="stSidebar"] div[role="radiogroup"] label{
  width: 100%;
  border: 1px solid rgba(255,255,255,0.22);
  background: rgba(255,255,255,0.06);
  border-radius: 14px;
  padding: 11px 12px;
  margin: 0px !important;
  cursor: pointer;
  transition: all 120ms ease-in-out;
}

section[data-testid="stSidebar"] div[role="radiogroup"] label:hover{
  background: rgba(255,255,255,0.14);
  border-color: rgba(255,255,255,0.38);
}

section[data-testid="stSidebar"] div[role="radiogroup"] label:has(input:checked){
  background: rgba(255,255,255,0.20);
  border-color: rgba(255,255,255,0.60);
}

section[data-testid="stSidebar"] div[role="radiogroup"]{
  gap: 10px;
}
</style>
""", unsafe_allow_html=True)

# =========================================================
# CSS: Blue sidebar + "button-like" radio
# =========================================================

/* Hover state */
section[data-testid="stSidebar"] div[role="radiogroup"] label:hover{
  background: rgba(255,255,255,0.14);
  border-color: rgba(255,255,255,0.38);
  transform: translateY(-1px);
}

/* Selected state */
section[data-testid="stSidebar"] div[role="radiogroup"] label:has(input:checked){
  background: rgba(255,255,255,0.20);
  border-color: rgba(255,255,255,0.60);
  box-shadow: 0 6px 18px rgba(0,0,0,0.18);
}

/* Spacing between buttons */
section[data-testid="stSidebar"] div[role="radiogroup"]{
  gap: 10px;
}

""", unsafe_allow_html=True)


# =========================================================
# Utilities: downloads
# =========================================================
def df_to_excel_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        for sheet, _df in sheets.items():
            _df.to_excel(writer, index=False, sheet_name=sheet[:31])
    bio.seek(0)
    return bio.getvalue()

def fig_to_png_bytes(fig: plt.Figure, dpi: int = 200) -> bytes:
    bio = io.BytesIO()
    fig.savefig(bio, format="png", dpi=dpi, bbox_inches="tight")
    bio.seek(0)
    return bio.getvalue()

def df_to_png_bytes(df: pd.DataFrame, title: str = "", dpi: int = 200) -> bytes:
    n_rows, n_cols = df.shape
    fig_w = min(18, max(6, n_cols * 1.5))
    fig_h = min(18, max(2.2, (n_rows + 1) * 0.45))
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")
    if title:
        ax.set_title(title, fontsize=12, pad=8)
    tbl = ax.table(cellText=df.values, colLabels=df.columns, cellLoc="center", loc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)
    tbl.scale(1, 1.15)
    png = fig_to_png_bytes(fig, dpi=dpi)
    plt.close(fig)
    return png

def download_table_block(df: pd.DataFrame, base_name: str, title: str = ""):
    c1, c2 = st.columns([1, 1])
    with c1:
        st.download_button(
            "Download Excel",
            data=df_to_excel_bytes({base_name: df}),
            file_name=f"{base_name}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    with c2:
        st.download_button(
            "Download PNG",
            data=df_to_png_bytes(df, title=title),
            file_name=f"{base_name}.png",
            mime="image/png",
            use_container_width=True
        )

def download_figure_block(fig: plt.Figure, base_name: str):
    st.download_button(
        "Download PNG",
        data=fig_to_png_bytes(fig),
        file_name=f"{base_name}.png",
        mime="image/png",
        use_container_width=False
    )


# =========================================================
# Template: generic (user chooses target/features)
# =========================================================
def make_template_excel_bytes() -> bytes:
    template = pd.DataFrame({
        "TargetBinary": [0, 1],
        "Feature1": [10.0, 50.0],
        "Feature2": [7.5, 14.2],
        "Feature3": [98, 92],
        "Feature4": [36.8, 39.1],
    })
    notes = pd.DataFrame({
        "Notes": [
            "This template is generic. Rename columns as needed.",
            "For Classification (Binary): target column must contain exactly two classes (e.g., 0/1).",
            "For Regression (Continuous): target column must be numeric with >2 unique values.",
            "Features should be numeric.",
            "You can add more feature columns (Feature5, Feature6, ...)."
        ]
    })
    return df_to_excel_bytes({"template": template, "notes": notes})


# =========================================================
# DeLong test (for ROC AUC difference)
# =========================================================
def compute_midrank(x):
    J = np.argsort(x)
    Z = x[J]
    N = len(x)
    T = np.zeros(N, dtype=float)
    i = 0
    while i < N:
        j = i
        while j < N and Z[j] == Z[i]:
            j += 1
        T[i:j] = 0.5 * (i + j - 1)
        i = j
    T2 = np.empty(N, dtype=float)
    T2[J] = T + 1
    return T2

def fastDeLong(predictions_sorted_transposed, label_1_count):
    m = label_1_count
    n = predictions_sorted_transposed.shape[1] - m
    pos = predictions_sorted_transposed[:, :m]
    neg = predictions_sorted_transposed[:, m:]
    k = predictions_sorted_transposed.shape[0]
    tx = np.empty([k, m])
    ty = np.empty([k, n])
    tz = np.empty([k, m + n])
    for r in range(k):
        tx[r] = compute_midrank(pos[r])
        ty[r] = compute_midrank(neg[r])
        tz[r] = compute_midrank(predictions_sorted_transposed[r])
    aucs = tx.sum(axis=1) / m / n - (m + 1) / (2 * n)
    v01 = (tz[:, :m] - tx) / n
    v10 = (tz[:, m:] - ty) / m
    sx = np.cov(v01)
    sy = np.cov(v10)
    return aucs, sx / m + sy / n

def delong_roc_test(y_true, y_scores_1, y_scores_2):
    y_true = np.array(y_true)
    order = np.argsort(-y_scores_1)
    y_true = y_true[order]
    y_scores_1 = y_scores_1[order]
    y_scores_2 = y_scores_2[order]
    label_1_count = int(np.sum(y_true))
    preds = np.vstack((y_scores_1, y_scores_2))
    aucs, cov = fastDeLong(preds, label_1_count)
    diff = aucs[0] - aucs[1]
    var = cov[0, 0] + cov[1, 1] - 2 * cov[0, 1]
    z = np.abs(diff) / np.sqrt(var)
    return float(2 * (1 - stats.norm.cdf(z)))


# =========================================================
# Data loading
# =========================================================
def load_dataframe(file):
    if file is None:
        return None
    name = file.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(file)
    return pd.read_excel(file)

def numeric_columns(df: pd.DataFrame) -> list[str]:
    return df.select_dtypes(include="number").columns.tolist()


# =========================================================
# Modeling: Classification (Binary)
# =========================================================
@st.cache_data(show_spinner=False)
def run_classification(df: pd.DataFrame, target: str, features: list[str],
                       test_size: float, random_state: int, n_bootstraps: int):
    data = df.copy()

    # Basic validation
    if target not in data.columns:
        raise ValueError("Target column not found.")
    for f in features:
        if f not in data.columns:
            raise ValueError(f"Feature column not found: {f}")

    # Drop rows with NA in used cols
    used = [target] + features
    data = data[used].dropna()

    if data[target].nunique() != 2:
        raise ValueError("Binary classification requires a target column with exactly 2 unique values (e.g., 0/1).")

    X = data[features]
    y = data[target].astype(int)

    # Statsmodels logistic (OR, CI, p-values)
    X_sm = sm.add_constant(X)
    logit = sm.Logit(y, X_sm).fit(disp=False)
    params = logit.params
    conf = logit.conf_int()
    pvals = logit.pvalues

    odds_table = pd.DataFrame({
        "Term": params.index,
        "Coefficient (β)": params.values,
        "Odds Ratio": np.exp(params.values),
        "CI 2.5%": np.exp(conf[0].values),
        "CI 97.5%": np.exp(conf[1].values),
        "p-value": pvals.values,
    }).round(4)
    odds_table["Significant (p<0.05)"] = odds_table["p-value"].apply(lambda p: "Yes" if p < 0.05 else "")

    # ML models
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=test_size, stratify=y, random_state=random_state
    )

    models = {
        "Logistic Regression": LogisticRegression(max_iter=2000),
        "Random Forest": RandomForestClassifier(),
        "SVM": SVC(probability=True),
        "k-NN": KNeighborsClassifier(),
    }

    probs = {}
    auc_scores = {}

    for name, clf in models.items():
        clf.fit(X_train, y_train)
        p = clf.predict_proba(X_test)[:, 1]
        probs[name] = p
        fpr, tpr, _ = roc_curve(y_test, p)
        auc_scores[name] = auc(fpr, tpr)

    # Bootstrap CI for AUC
    def bootstrap_auc_ci(y_true, y_scores, n_bootstraps=200, alpha=0.05):
        rng = np.random.RandomState(42)
        y_true = np.asarray(y_true)
        y_scores = np.asarray(y_scores)
        scores = []
        for _ in range(n_bootstraps):
            idx = rng.randint(0, len(y_scores), len(y_scores))
            if len(np.unique(y_true[idx])) < 2:
                continue
            fpr, tpr, _ = roc_curve(y_true[idx], y_scores[idx])
            scores.append(auc(fpr, tpr))
        scores = np.array(scores)
        scores.sort()
        lo = scores[int((alpha / 2) * len(scores))]
        hi = scores[int((1 - alpha / 2) * len(scores))]
        return float(lo), float(hi)

    rows = []
    for m in auc_scores:
        lo, hi = bootstrap_auc_ci(y_test.values, probs[m], n_bootstraps=n_bootstraps)
        rows.append({
            "Model": m,
            "AUC": round(auc_scores[m], 4),
            "95% CI Lower": round(lo, 4),
            "95% CI Upper": round(hi, 4),
        })
    auc_df = pd.DataFrame(rows).sort_values("AUC", ascending=False).reset_index(drop=True)

    top1, top2 = auc_df["Model"].iloc[0], auc_df["Model"].iloc[1]
    delong_p = delong_roc_test(y_test.values, probs[top1], probs[top2])

    # ROC figure
    roc_fig = plt.figure(figsize=(8, 5.5))
    for name in models:
        fpr, tpr, _ = roc_curve(y_test, probs[name])
        plt.plot(fpr, tpr, label=f"{name} (AUC={auc_scores[name]:.3f})", linewidth=2)
    plt.plot([0, 1], [0, 1], linestyle="--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("ROC Curve")
    plt.grid(True)
    plt.legend(loc="lower right")
    plt.tight_layout()

    # Reports + Confusion matrices (top2)
    reports = {}
    cm_figs = {}
    for m in [top1, top2]:
        y_pred = (probs[m] >= 0.5).astype(int)
        rep = pd.DataFrame(classification_report(y_test, y_pred, output_dict=True)).T.round(4)
        reports[m] = rep

        cm = confusion_matrix(y_test, y_pred)
        fig = plt.figure(figsize=(5.2, 4.2))
        disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=["No", "Yes"])
        disp.plot(values_format="d")
        plt.title(f"Confusion Matrix — {m}")
        plt.tight_layout()
        cm_figs[m] = fig

    return {
        "used_data": data,
        "odds_table": odds_table,
        "auc_df": auc_df,
        "roc_fig": roc_fig,
        "top1": top1,
        "top2": top2,
        "delong_p": float(delong_p),
        "reports": reports,
        "cm_figs": cm_figs,
        "target": target,
        "features": features,
    }


# =========================================================
# Modeling: Regression (Continuous) — OLS + optional Ridge/Lasso
# =========================================================
@st.cache_data(show_spinner=False)
def run_regression(df: pd.DataFrame, target: str, features: list[str],
                   test_size: float, random_state: int, reg_model: str):
    data = df.copy()

    if target not in data.columns:
        raise ValueError("Target column not found.")
    for f in features:
        if f not in data.columns:
            raise ValueError(f"Feature column not found: {f}")

    used = [target] + features
    data = data[used].dropna()

    if data[target].nunique() <= 2:
        raise ValueError("Continuous regression requires a numeric target with more than 2 unique values.")

    X = data[features]
    y = data[target].astype(float)

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=test_size, random_state=random_state
    )

    # Statsmodels OLS for coefficients, CI, p-values
    X_sm = sm.add_constant(X)
    ols = sm.OLS(y, X_sm).fit()

    params = ols.params
    conf = ols.conf_int()
    pvals = ols.pvalues

    coef_table = pd.DataFrame({
        "Term": params.index,
        "Coefficient (β)": params.values,
        "CI 2.5%": conf[0].values,
        "CI 97.5%": conf[1].values,
        "p-value": pvals.values,
    }).round(4)
    coef_table["Significant (p<0.05)"] = coef_table["p-value"].apply(lambda p: "Yes" if p < 0.05 else "")

    # Predictive model (optional): LR/Ridge/Lasso
    if reg_model == "LinearRegression":
        model = LinearRegression()
    elif reg_model == "Ridge":
        model = Ridge(alpha=1.0)
    else:
        model = Lasso(alpha=0.01, max_iter=5000)

    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)

    mae = float(mean_absolute_error(y_test, y_pred))
    rmse = float(np.sqrt(mean_squared_error(y_test, y_pred)))
    r2 = float(r2_score(y_test, y_pred))

    metrics = pd.DataFrame([{
        "Model": reg_model,
        "MAE": round(mae, 4),
        "RMSE": round(rmse, 4),
        "R²": round(r2, 4),
        "OLS Adj. R²": round(float(ols.rsquared_adj), 4),
    }])

    # Figures
    fig_pred = plt.figure(figsize=(6.5, 4.8))
    ax = fig_pred.add_subplot(111)
    ax.scatter(y_test, y_pred)
    ax.set_title("Predicted vs Actual")
    ax.set_xlabel("Actual")
    ax.set_ylabel("Predicted")
    ax.grid(True)
    plt.tight_layout()

    residuals = y_test.values - y_pred
    fig_res = plt.figure(figsize=(6.5, 4.8))
    ax2 = fig_res.add_subplot(111)
    ax2.scatter(y_pred, residuals)
    ax2.axhline(0, linestyle="--")
    ax2.set_title("Residuals vs Predicted")
    ax2.set_xlabel("Predicted")
    ax2.set_ylabel("Residuals")
    ax2.grid(True)
    plt.tight_layout()

    return {
        "used_data": data,
        "coef_table": coef_table,
        "metrics": metrics,
        "fig_pred": fig_pred,
        "fig_res": fig_res,
        "target": target,
        "features": features,
        "reg_model": reg_model
    }


# =========================================================
# Export helpers
# =========================================================
def build_docx_classification(result: dict) -> bytes:
    doc = Document()
    doc.add_heading("Binary Classification Report", level=1)
    doc.add_paragraph(f"Target: {result['target']}")
    doc.add_paragraph(f"Features: {', '.join(result['features'])}")

    doc.add_heading("AUC Summary", level=2)
    doc.add_paragraph(result["auc_df"].to_string(index=False))

    doc.add_heading("Logistic Regression (Odds Ratios)", level=2)
    doc.add_paragraph(result["odds_table"].to_string(index=False))

    doc.add_heading("DeLong Test", level=2)
    doc.add_paragraph(f"Top1: {result['top1']} | Top2: {result['top2']} | p-value: {result['delong_p']:.6f}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def build_docx_regression(result: dict) -> bytes:
    doc = Document()
    doc.add_heading("Continuous Regression Report", level=1)
    doc.add_paragraph(f"Target: {result['target']}")
    doc.add_paragraph(f"Features: {', '.join(result['features'])}")
    doc.add_paragraph(f"Predictive model: {result['reg_model']}")

    doc.add_heading("Metrics", level=2)
    doc.add_paragraph(result["metrics"].to_string(index=False))

    doc.add_heading("OLS Coefficients", level=2)
    doc.add_paragraph(result["coef_table"].to_string(index=False))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# =========================================================
# Sidebar: modules + data + settings
# =========================================================
with st.sidebar:
    module = st.radio(
        "",
        ["Data", "Explore (EDA)", "Modeling", "Reports", "Export"],
        index=0,
        label_visibility="collapsed"
    )

    st.divider()

    st.download_button(
        "Download Excel template",
        data=make_template_excel_bytes(),
        file_name="logistic_template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )
    uploaded = st.file_uploader("Upload CSV/XLSX", type=["csv", "xlsx"])

    with st.expander("Settings", expanded=False):
        problem_type = st.radio(
            "Problem type",
            ["Classification (Binary)", "Regression (Continuous)"],
            index=0
        )
        test_size = st.slider("Test size", 0.1, 0.5, 0.3, 0.05)
        random_state = st.number_input("Random state", value=42, step=1)

        # Keep bootstrap lighter for web
        n_bootstraps = st.slider("Bootstrap iterations (AUC CI)", 100, 500, 200, 50)

        reg_model = st.selectbox("Regression model (for prediction)", ["LinearRegression", "Ridge", "Lasso"])

    st.caption("Upload CSV/XLSX • Use the template for best results.")


# =========================================================
# Main: Title (no icon card, no extra text)
# =========================================================
st.title("Disease Diagnosis with Logistic Regression")
# If you want English title instead, use:
# st.title("Disease Diagnosis with Logistic Regression")


df = load_dataframe(uploaded)


# =========================================================
# Data page
# =========================================================
if module == "Data":
    st.subheader("Data")
    if df is None:
        st.info("Upload a CSV/XLSX file from the sidebar to get started.")
        st.stop()

    left, right = st.columns([2, 1])
    with left:
        st.markdown("#### Preview")
        st.dataframe(df.head(30), use_container_width=True)
        download_table_block(df.head(200).reset_index(drop=True), "data_preview", "Data Preview (first rows)")

    with right:
        st.markdown("#### Quick checks")
        numeric_cols = numeric_columns(df)
        info = pd.DataFrame({
            "Rows": [df.shape[0]],
            "Columns": [df.shape[1]],
            "Numeric columns": [len(numeric_cols)],
            "Missing values": [int(df.isna().sum().sum())],
            "Duplicate rows": [int(df.duplicated().sum())],
        })
        st.dataframe(info, use_container_width=True)
        download_table_block(info, "data_quick_checks", "Quick Checks")

        st.markdown("#### Column types")
        types_df = pd.DataFrame({"Column": df.columns, "Dtype": df.dtypes.astype(str).values})
        st.dataframe(types_df, use_container_width=True, height=280)
        download_table_block(types_df, "data_column_types", "Column Types")


# =========================================================
# EDA
# =========================================================
elif module == "Explore (EDA)":
    st.subheader("Explore (EDA)")
    if df is None:
        st.info("Upload a CSV/XLSX file from the sidebar to get started.")
        st.stop()

    num_cols = numeric_columns(df)
    if not num_cols:
        st.warning("No numeric columns found. EDA requires numeric data.")
        st.stop()

    tab1, tab2, tab3 = st.tabs(["Summary", "Distributions", "Correlations"])

    with tab1:
        st.markdown("#### Descriptive statistics (numeric columns)")
        desc = df[num_cols].describe().T.round(4).reset_index(names="Feature")
        st.dataframe(desc, use_container_width=True)
        download_table_block(desc, "eda_descriptive_stats", "Descriptive Statistics")

        st.markdown("#### Missing values by column")
        miss = df.isna().sum().reset_index()
        miss.columns = ["Column", "Missing"]
        st.dataframe(miss, use_container_width=True)
        download_table_block(miss, "eda_missing_by_column", "Missing Values by Column")

    with tab2:
        feature = st.selectbox("Select a numeric feature", num_cols)
        fig = plt.figure(figsize=(7.0, 4.6))
        ax = fig.add_subplot(111)
        sns.histplot(data=df, x=feature, kde=True, ax=ax)
        ax.set_title(f"Distribution — {feature}")
        ax.grid(True)
        plt.tight_layout()
        st.pyplot(fig, clear_figure=False)
        download_figure_block(fig, f"eda_hist_{feature}".lower())

    with tab3:
        if len(num_cols) < 2:
            st.warning("Need at least 2 numeric columns for correlation.")
            st.stop()

        corr = df[num_cols].corr(numeric_only=True).round(4)
        fig = plt.figure(figsize=(7.5, 5.5))
        ax = fig.add_subplot(111)
        sns.heatmap(corr, annot=False, ax=ax)
        ax.set_title("Correlation heatmap")
        plt.tight_layout()
        st.pyplot(fig, clear_figure=False)
        download_figure_block(fig, "eda_correlation_heatmap")

        corr_table = corr.reset_index(names="Feature")
        st.dataframe(corr_table, use_container_width=True)
        download_table_block(corr_table, "eda_correlation_table", "Correlation Table")


# =========================================================
# Modeling (generic)
# =========================================================
elif module == "Modeling":
    st.subheader("Modeling")
    if df is None:
        st.info("Upload a CSV/XLSX file from the sidebar to get started.")
        st.stop()

    num_cols = numeric_columns(df)
    if not num_cols:
        st.warning("No numeric columns found. Modeling requires numeric data.")
        st.stop()

    # Candidate targets depending on problem type
    if problem_type == "Classification (Binary)":
        candidate_targets = [c for c in num_cols if df[c].dropna().nunique() == 2]
        help_target = "Select a binary target column (must contain exactly 2 unique values, e.g., 0/1)."
    else:
        candidate_targets = [c for c in num_cols if df[c].dropna().nunique() > 2]
        help_target = "Select a continuous numeric target column (>2 unique values)."

    if not candidate_targets:
        st.error("No suitable target columns found for the selected problem type.")
        st.stop()

    target = st.selectbox("Target column", candidate_targets, help=help_target)
    feature_candidates = [c for c in num_cols if c != target]
    default_feats = feature_candidates[: min(4, len(feature_candidates))]
    features = st.multiselect(
        "Feature columns",
        feature_candidates,
        default=default_feats,
        help="Select numeric feature columns used by the model."
    )

    if len(features) == 0:
        st.warning("Select at least one feature column.")
        st.stop()

    run = st.button("Run analysis", type="primary", use_container_width=True)
    if not run:
        st.info("Click **Run analysis** to compute results.")
        st.stop()

    if problem_type == "Classification (Binary)":
        with st.spinner("Running binary classification..."):
            cls_result = run_classification(
                df=df,
                target=target,
                features=features,
                test_size=float(test_size),
                random_state=int(random_state),
                n_bootstraps=int(n_bootstraps),
            )

        st.markdown("### Logistic regression (Odds Ratios)")
        st.dataframe(cls_result["odds_table"], use_container_width=True)
        download_table_block(cls_result["odds_table"], "logistic_odds_ratios", "Odds Ratios")

        st.markdown("### Model comparison (AUC with 95% CI)")
        st.dataframe(cls_result["auc_df"], use_container_width=True)
        download_table_block(cls_result["auc_df"], "auc_summary", "AUC Summary")

        st.markdown("### ROC curve")
        st.pyplot(cls_result["roc_fig"], clear_figure=False)
        download_figure_block(cls_result["roc_fig"], "roc_curve")

        st.markdown("### DeLong test (Top 2 models)")
        delong_df = pd.DataFrame([{
            "Top 1": cls_result["top1"],
            "Top 2": cls_result["top2"],
            "p-value": round(cls_result["delong_p"], 6),
        }])
        st.dataframe(delong_df, use_container_width=True)
        download_table_block(delong_df, "delong_test", "DeLong Test")

        st.markdown("### Top 2 model reports")
        for m in [cls_result["top1"], cls_result["top2"]]:
            st.markdown(f"#### {m}")
            c1, c2 = st.columns([1.2, 1])
            with c1:
                rep = cls_result["reports"][m].reset_index(names="Metric")
                st.dataframe(rep, use_container_width=True)
                download_table_block(rep, f"classification_report_{m}".replace(" ", "_").lower(), f"Classification Report — {m}")
            with c2:
                fig_cm = cls_result["cm_figs"][m]
                st.pyplot(fig_cm, clear_figure=False)
                download_figure_block(fig_cm, f"confusion_matrix_{m}".replace(" ", "_").lower())

        st.session_state["last_result"] = {"type": "classification", "payload": cls_result}

    else:
        with st.spinner("Running continuous regression..."):
            reg_result = run_regression(
                df=df,
                target=target,
                features=features,
                test_size=float(test_size),
                random_state=int(random_state),
                reg_model=str(reg_model)
            )

        st.markdown("### OLS coefficients (inference)")
        st.dataframe(reg_result["coef_table"], use_container_width=True)
        download_table_block(reg_result["coef_table"], "ols_coefficients", "OLS Coefficients")

        st.markdown("### Predictive performance")
        st.dataframe(reg_result["metrics"], use_container_width=True)
        download_table_block(reg_result["metrics"], "regression_metrics", "Regression Metrics")

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("### Predicted vs Actual")
            st.pyplot(reg_result["fig_pred"], clear_figure=False)
            download_figure_block(reg_result["fig_pred"], "predicted_vs_actual")
        with c2:
            st.markdown("### Residuals vs Predicted")
            st.pyplot(reg_result["fig_res"], clear_figure=False)
            download_figure_block(reg_result["fig_res"], "residuals_vs_predicted")

        st.session_state["last_result"] = {"type": "regression", "payload": reg_result}


# =========================================================
# Reports
# =========================================================
elif module == "Reports":
    st.subheader("Reports")
    last = st.session_state.get("last_result")

    if not last:
        st.info("Run an analysis in the **Modeling** module first.")
        st.stop()

    if last["type"] == "classification":
        r = last["payload"]
        st.markdown("#### Binary classification summary")
        summary = pd.DataFrame([{
            "Target": r["target"],
            "Features": ", ".join(r["features"]),
            "Top 1 model": r["top1"],
            "Top 2 model": r["top2"],
            "DeLong p-value": round(r["delong_p"], 6),
        }])
        st.dataframe(summary, use_container_width=True)
        download_table_block(summary, "classification_summary", "Classification Summary")

    else:
        r = last["payload"]
        st.markdown("#### Continuous regression summary")
        summary = pd.DataFrame([{
            "Target": r["target"],
            "Features": ", ".join(r["features"]),
            "Predictive model": r["reg_model"],
            "MAE": float(r["metrics"]["MAE"].iloc[0]),
            "RMSE": float(r["metrics"]["RMSE"].iloc[0]),
            "R²": float(r["metrics"]["R²"].iloc[0]),
        }])
        st.dataframe(summary, use_container_width=True)
        download_table_block(summary, "regression_summary", "Regression Summary")


# =========================================================
# Export (Excel + Word)
# =========================================================
elif module == "Export":
    st.subheader("Export")
    last = st.session_state.get("last_result")

    if df is None:
        st.info("Upload a dataset first.")
        st.stop()

    if not last:
        st.info("Run an analysis in the **Modeling** module first.")
        st.stop()

    # Always allow exporting data
    st.markdown("#### Dataset export")
    st.download_button(
        "Download dataset (Excel)",
        data=df_to_excel_bytes({"data": df}),
        file_name="dataset.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

    st.markdown("#### Analysis export")
    if last["type"] == "classification":
        r = last["payload"]
        excel_bytes = df_to_excel_bytes({
            "used_data": r["used_data"],
            "odds_ratios": r["odds_table"],
            "auc_summary": r["auc_df"],
            "delong": pd.DataFrame([{"Top 1": r["top1"], "Top 2": r["top2"], "p-value": r["delong_p"]}]),
        })
        st.download_button(
            "Download analysis (Excel)",
            data=excel_bytes,
            file_name="classification_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

        docx_bytes = build_docx_classification(r)
        st.download_button(
            "Download analysis (Word)",
            data=docx_bytes,
            file_name="classification_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    else:
        r = last["payload"]
        excel_bytes = df_to_excel_bytes({
            "used_data": r["used_data"],
            "metrics": r["metrics"],
            "ols_coefficients": r["coef_table"],
        })
        st.download_button(
            "Download analysis (Excel)",
            data=excel_bytes,
            file_name="regression_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

        docx_bytes = build_docx_regression(r)
        st.download_button(
            "Download analysis (Word)",
            data=docx_bytes,
            file_name="regression_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
